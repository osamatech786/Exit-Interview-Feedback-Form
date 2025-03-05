import streamlit as st
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
import os
from datetime import datetime  # Added for timestamp
import time  # Added for timeout handling

# Set the page configuration for the Streamlit app
st.set_page_config(
    page_title="Exit Interview Feedback Form", 
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"
)

if 'submission_status' not in st.session_state: 
    st.session_state.submission_status = False

# Function to replace placeholders in the Word document (robust and simple, preserving brackets, treating placeholders as complete words)
def replace_placeholder(doc, placeholder, value):
    """ Replace the placeholder text (e.g., p1, [p1]) with the value, preserving brackets if present, treating placeholders as whole words. """
    placeholder_text = f'{placeholder}'  # Use just the placeholder (e.g., p1, p10)
    value_str = str(value) if value is not None else ""

    # Process table cells first (prioritize rows, simple traversal)
    processed_tables = set()  # Track processed tables to avoid duplicates
    for table in doc.tables:
        if id(table) in processed_tables:
            continue  # Skip already processed tables
        processed_tables.add(id(table))
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
                    if (f'[{placeholder_text}]' in cell.text or 
                        placeholder_text in cell.text.split()):
                        # Preserve brackets if present, replace placeholder as a whole word
                        cell.text = cell.text.replace(f'[{placeholder_text}]', f'[{value_str}]').replace(placeholder_text, value_str)

    # Process paragraphs and their runs (for non-table content)
    for element in doc.paragraphs:
        if hasattr(element, 'text') and element.text:
            # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
            if (f'[{placeholder_text}]' in element.text or 
                placeholder_text in element.text.split()):
                # Preserve brackets if present, replace placeholder as a whole word
                element.text = element.text.replace(f'[{placeholder_text}]', f'[{value_str}]').replace(placeholder_text, value_str)
        if hasattr(element, 'runs'):
            for run in element.runs:
                if run.text:
                    # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
                    if (f'[{placeholder_text}]' in run.text or 
                        placeholder_text in run.text.split()):
                        # Preserve brackets if present, replace placeholder as a whole word
                        run.text = run.text.replace(f'[{placeholder_text}]', f'[{value_str}]').replace(placeholder_text, value_str)

# Function to mark selected options with an 'X' in the Word document (robust and simple, preserving brackets, treating placeholders as complete words)
def mark_selected_option(doc, placeholder_dict):
    """ Mark the selected option with an 'X' or space, preserving brackets if present, treating placeholders as whole words. """
    for placeholder, is_selected in placeholder_dict.items():
        placeholder_text = f'{placeholder}'  # Use just the placeholder (e.g., p1, p10)
        marker = 'X' if is_selected else ' '  # Use 'X' for selected, space for unselected

        # Process table cells first (prioritize rows, simple traversal)
        processed_tables = set()  # Track processed tables to avoid duplicates
        for table in doc.tables:
            if id(table) in processed_tables:
                continue  # Skip already processed tables
            processed_tables.add(id(table))
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
                        if (f'[{placeholder_text}]' in cell.text or 
                            placeholder_text in cell.text.split()):
                            # Preserve brackets if present, replace placeholder as a whole word
                            cell.text = cell.text.replace(f'[{placeholder_text}]', f'[{marker}]').replace(placeholder_text, marker)

        # Process paragraphs and their runs (for non-table content)
        for element in doc.paragraphs:
            if hasattr(element, 'text') and element.text:
                # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
                if (f'[{placeholder_text}]' in element.text or 
                    placeholder_text in element.text.split()):
                    # Preserve brackets if present, replace placeholder as a whole word
                    element.text = element.text.replace(f'[{placeholder_text}]', f'[{marker}]').replace(placeholder_text, marker)
            if hasattr(element, 'runs'):
                for run in element.runs:
                    if run.text:
                        # Use simple word boundaries to ensure p1 and p10 are treated as separate placeholders
                        if (f'[{placeholder_text}]' in run.text or 
                            placeholder_text in run.text.split()):
                            # Preserve brackets if present, replace placeholder as a whole word
                            run.text = run.text.replace(f'[{placeholder_text}]', f'[{marker}]').replace(placeholder_text, marker)

# Function to populate the Word document with form data (robust and simple)
def populate_document(data, template_path, save_directory="/"):
    try:
        if not os.path.exists(save_directory):
            os.makedirs(save_directory, exist_ok=True)
        
        doc = Document(template_path)

        # Employee Information (plain placeholders, e.g., p21)
        replace_placeholder(doc, 'p21', data['name'])
        replace_placeholder(doc, 'p22', data['department'])
        replace_placeholder(doc, 'p23', data['job_title'])
        replace_placeholder(doc, 'p24', data['last_working_day'])

        # 1. Reason for Leaving (placeholders, e.g., p1)
        mark_selected_option(doc, {
            'p1': data['reason_for_leaving'] == "New Job Opportunity",
            'p2': data['reason_for_leaving'] == "Career Growth",
            'p3': data['reason_for_leaving'] == "Work-Life Balance",
            'p4': data['reason_for_leaving'] == "Salary & Benefits",
            'p5': data['reason_for_leaving'] == "Work Environment",
            'p6': data['reason_for_leaving'] == "Management Issues",
            'p7': data['reason_for_leaving'] == "Personal Reasons",
            'p8': data['reason_for_leaving'] == "Other"
        })
        replace_placeholder(doc, 'p9', data['other_reason'] if data['reason_for_leaving'] == "Other" else "")

        # 2. Job Experience & Work Environment (placeholders, e.g., p10)
        replace_placeholder(doc, 'p10', data['enjoyed_most'])
        replace_placeholder(doc, 'p11', data['challenges'])
        replace_placeholder(doc, 'p12', str(data['manager_relationship']))
        mark_selected_option(doc, {
            'p13': data['training_opportunities'] == "Yes",
            'p14': data['training_opportunities'] == "No"
        })

        # 3. Compensation & Benefits (placeholders, e.g., p15)
        replace_placeholder(doc, 'p15', str(data['salary_satisfaction']))
        replace_placeholder(doc, 'p16', str(data['benefits_satisfaction']))

        # 4. Suggestions for Improvement (placeholders, e.g., p17)
        replace_placeholder(doc, 'p17', data['recommendations'])
        mark_selected_option(doc, {
            'p18': data['recommend_company'] == "Yes",
            'p19': data['recommend_company'] == "No"
        })

        # 5. Final Comments (placeholder, e.g., p20)
        replace_placeholder(doc, 'p20', data['final_comments'])

        # Save with name and timestamp to avoid overwrites
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filled_doc_path = f"Exit_Interview_Form_{data['name']}_{timestamp}.docx"
        doc.save(filled_doc_path)

        return filled_doc_path

    except Exception as e:
        st.error(f"Error processing the document: {e}")
        raise  # Re-raise to ensure the spinner stops

# Function to send the document via Outlook with timeout (simple and robust)
def send_email(file_path, timeout=30):
    try:
        start_time = time.time()
        sender_email = st.secrets["sender_email"]
        password = st.secrets["sender_password"]
        receiver_email = sender_email
        smtp_server = "smtp.office365.com"
        smtp_port = 587

        if not os.path.exists(file_path):
            st.warning("File not found. Skipping email sending.")
            return False

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = "Exit Interview Feedback Form Submission"
        body = "Please find the attached filled exit interview feedback form."
        msg.attach(MIMEText(body, 'plain'))

        with open(file_path, "rb") as attachment:
            part = MIMEApplication(attachment.read(), _subtype="docx")
            part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(file_path)}"')
            msg.attach(part)

        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()

        return True

    except smtplib.SMTPException as smtp_error:
        st.error(f"SMTP error occurred: {smtp_error}")
        return False
    except Exception as e:
        st.error(f"An error occurred while sending the email: {e}")
        return False
    finally:
        if time.time() - start_time > timeout:
            st.error("Email sending timed out. Please try again or check your SMTP settings.")
            return False

# Streamlit form
st.title("Exit Interview Feedback Form")
st.write("________________________________________")
st.write("Thank you for taking the time to provide feedback on your experience. Your input is valuable and will help us improve. Please fill out the form below.")

st.subheader("Employee Information")
st.text_input("Name", key="name")
st.text_input("Department", key="department")
st.text_input("Job Title", key="job_title")
st.date_input("Last Working Day", key="last_working_day", format='DD/MM/YYYY')

st.subheader("1. Reason for Leaving")
st.radio(
    "Please select the primary reason for your departure:",
    ["New Job Opportunity", "Career Growth", "Work-Life Balance", "Salary & Benefits", 
     "Work Environment", "Management Issues", "Personal Reasons", "Other"],
    key="reason_for_leaving"
)
st.text_input("If 'Other', please specify:", key="other_reason", disabled=st.session_state.reason_for_leaving != "Other")

st.subheader("2. Job Experience & Work Environment")
st.text_area("What aspects of your job did you enjoy the most?", key="enjoyed_most")
st.text_area("What aspects did you find challenging or frustrating?", key="challenges")
st.slider("How would you rate your relationship with your manager? (1-5)", 1, 5, 3, key="manager_relationship")
st.radio("Were you provided with enough training and development opportunities?", ["Yes", "No"], key="training_opportunities")

st.subheader("3. Compensation & Benefits")
st.slider("How satisfied were you with your salary? (1-5)", 1, 5, 3, key="salary_satisfaction")
st.slider("How satisfied were you with the benefits package? (1-5)", 1, 5, 3, key="benefits_satisfaction")

st.subheader("4. Suggestions for Improvement")
st.text_area("What changes would you recommend for the company?", key="recommendations")
st.radio("Would you recommend this company to others as a good workplace?", ["Yes", "No"], key="recommend_company")

st.subheader("5. Final Comments")
st.text_area("Any additional feedback:", key="final_comments")

# Submit button with timeout
if st.button("Submit", key="submit_button", disabled=st.session_state.submission_status):
    with st.spinner('Processing...'):
        try:
            start_time = time.time()
            timeout = 60  # Set a 60-second timeout for the entire process

            if not st.session_state.name or not st.session_state.department or not st.session_state.job_title:
                st.error("Please fill in all required fields (Name, Department, Job Title).")
                st.session_state.submission_status = False
            else:
                form_data = {
                    'name': st.session_state.name,
                    'department': st.session_state.department,
                    'job_title': st.session_state.job_title,
                    'last_working_day': st.session_state.last_working_day.strftime("%d-%m-%Y"),
                    'reason_for_leaving': st.session_state.reason_for_leaving,
                    'other_reason': st.session_state.other_reason,
                    'enjoyed_most': st.session_state.enjoyed_most,
                    'challenges': st.session_state.challenges,
                    'manager_relationship': st.session_state.manager_relationship,
                    'training_opportunities': st.session_state.training_opportunities,
                    'salary_satisfaction': st.session_state.salary_satisfaction,
                    'benefits_satisfaction': st.session_state.benefits_satisfaction,
                    'recommendations': st.session_state.recommendations,
                    'recommend_company': st.session_state.recommend_company,
                    'final_comments': st.session_state.final_comments,
                }

                template_path = "resource/ph_Exit_Interview_Feedback_Form.docx"
                filled_doc_path = populate_document(form_data, template_path)

                if filled_doc_path:
                    if send_email(filled_doc_path, timeout=30):
                        st.session_state.submission_status = True
                        try:
                            with open(filled_doc_path, 'rb') as f:
                                file_contents = f.read()
                                st.download_button(
                                    label="Download Your Response",
                                    data=file_contents,
                                    file_name=filled_doc_path,
                                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                                )
                        except FileNotFoundError as file_error:
                            st.error(f"Error with file handling: {file_error}")
                            st.session_state.submission_status = False
                    else:
                        st.error("Failed to send email. Please check your SMTP settings or try again.")
                        st.session_state.submission_status = False
                else:
                    st.error("Failed to generate document. Please check the template path or document structure.")
                    st.session_state.submission_status = False

            if time.time() - start_time > timeout:
                st.error("Processing timed out after 60 seconds. Please check the document or network issues.")
                st.session_state.submission_status = False

        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            st.session_state.submission_status = False
        finally:
            if st.session_state.submission_status:
                st.success(f"Feedback form submitted successfully.")
            else:
                st.stop()  # Force the spinner to stop if submission fails

# if st.session_state.submission_status:
#     st.success(f"Feedback form submitted and sent to {st.secrets['sender_email']}.")

# streamlit run app.py
# Dev: https://linkedin.com/in/osamatech786